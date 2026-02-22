#!/usr/bin/env python3
"""
apa7_format.py
==============
Post-processes a Pandoc-generated DOCX to enforce APA7 formatting.

Usage:
    python3 apa7_format.py input.docx output.docx

What this script applies automatically
---------------------------------------
  - Times New Roman 12pt throughout (body, tables, captions)
  - Double line spacing with no extra space before/after paragraphs
  - 1-inch margins on all sides
  - 0.5-inch first-line indent for body paragraphs
  - No first-line indent on: abstract, first paragraph after a heading,
    table/figure labels, table titles, *Note.* lines, Keywords line
  - Hanging indent (0.5 in left, -0.5 in first-line) for reference entries
  - APA7 table borders: top, below-header row, and bottom only — no verticals
  - Heading Level 1: centered, bold
  - Heading Level 2: left-aligned, bold
  - Heading Level 3: left-aligned, bold italic
  - Code/verbatim blocks: left unchanged

What still needs manual attention in Word
------------------------------------------
  1. Page numbers  — Insert > Page Number > Top of Page > Plain Number 3
  2. Title page    — name, course, instructor, institution, date
  3. Figure labels — change *italic* "Figure N" to **bold** (the script cannot
                     distinguish label runs from caption runs reliably)
  4. *Note.* lines — confirm only the word "Note." is italic, not the full line
"""

import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

FONT_NAME = "Times New Roman"
FONT_SIZE = 12  # pt

# Paragraph text patterns that should never receive a first-line indent.
# Covers:  Table N / Figure N labels,  *Note.* lines,  Keywords: line
NO_INDENT_RE = re.compile(
    r"^(Table|Figure)\s+\d+|^Note\.|^Keywords?:",
    re.IGNORECASE,
)

# Style name substrings that identify code/verbatim blocks (leave unchanged)
CODE_STYLE_KEYWORDS = ("code", "verbatim", "source")


# ---------------------------------------------------------------------------
# Font helpers
# ---------------------------------------------------------------------------

def _set_rpr_defaults(para, name=FONT_NAME, size=FONT_SIZE):
    """
    Inject paragraph-level rPr defaults inside pPr so that any run
    added later (or runs with no explicit font set) inherits the right font.
    """
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)

    # sz values are in half-points
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), str(size * 2))


def apply_font(para, name=FONT_NAME, size=FONT_SIZE):
    """Apply font to every run in a paragraph and set paragraph-level defaults."""
    for run in para.runs:
        run.font.name = name
        run.font.size = Pt(size)
    _set_rpr_defaults(para, name, size)


# ---------------------------------------------------------------------------
# Spacing helpers
# ---------------------------------------------------------------------------

def set_double_spacing(para):
    """Double line spacing, no extra space before or after."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Heading helpers
# ---------------------------------------------------------------------------

def format_heading(para, level):
    """
    Apply APA7 heading formatting.
      Level 1 — centered, bold
      Level 2 — left-aligned, bold
      Level 3 — left-aligned, bold italic
    """
    apply_font(para)
    set_double_spacing(para)
    para.paragraph_format.first_line_indent = Pt(0)

    if level == 1:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold = True
            run.font.italic = False

    elif level == 2:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.bold = True
            run.font.italic = False

    elif level == 3:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.bold = True
            run.font.italic = True


# ---------------------------------------------------------------------------
# Table border helpers
# ---------------------------------------------------------------------------

def _border_element(side, val="none", sz="0", color="auto"):
    """Create a single border XML element."""
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def _clear_table_level_borders(table):
    """Remove all borders from the tblBorders element."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tblBorders.append(_border_element(side))
    tblPr.append(tblBorders)


def _clear_cell_borders(cell):
    """Remove all borders from a single table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)

    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tcBorders.append(_border_element(side))
    tcPr.append(tcBorders)


def _add_cell_borders(cell, sides, sz="8", color="000000"):
    """
    Add single-line borders to the specified sides of a cell.
    sz is in eighths of a point (8 = 1pt line).
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for side in sides:
        el = tcBorders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def remove_bookmarks(doc):
    """
    Remove all bookmark elements inserted by Pandoc.
    Pandoc adds w:bookmarkStart / w:bookmarkEnd to every heading for
    internal linking — these are unnecessary in a submitted paper and
    can cause unexpected behaviour in Word.
    """
    body = doc.element.body
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for el in body.findall(".//" + qn(tag)):
            el.getparent().remove(el)


def format_apa7_table(table):
    """
    Apply APA7 table border rules:
      - Horizontal line above the header row  (top of table)
      - Horizontal line below the header row  (header / data separator)
      - Horizontal line below the last row    (bottom of table)
      - No vertical lines anywhere
      - No horizontal lines between data rows
    Also sets Times New Roman 12pt, single spacing inside cells.
    """
    # 1. Strip all table-level borders
    _clear_table_level_borders(table)

    # 2. Strip all cell-level borders
    for row in table.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)

    # 3. Add the three required horizontal lines
    for cell in table.rows[0].cells:       # top of table
        _add_cell_borders(cell, ["top"])
    for cell in table.rows[0].cells:       # below header row
        _add_cell_borders(cell, ["bottom"])
    for cell in table.rows[-1].cells:      # bottom of table
        _add_cell_borders(cell, ["bottom"])

    # 4. Format text inside cells
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                apply_font(para)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para.paragraph_format.line_spacing = 1.0


# ---------------------------------------------------------------------------
# Main formatting function
# ---------------------------------------------------------------------------

def format_document(input_path, output_path):
    doc = Document(input_path)

    # ── 1. Margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── 2. Paragraphs ─────────────────────────────────────────────────────
    # State tracking
    section_name      = "body"   # "abstract" | "body" | "references"
    prev_was_heading  = False
    prev_was_label    = False    # True after a Table N / Figure N label line

    for para in doc.paragraphs:
        style = para.style.name
        text  = para.text.strip()

        # Leave code/verbatim blocks completely unchanged
        if any(kw in style.lower() for kw in CODE_STYLE_KEYWORDS):
            prev_was_heading = False
            prev_was_label   = False
            continue

        # ── Track which major section we're in ──
        if style.startswith("Heading"):
            tl = text.lower()
            if tl == "abstract":
                section_name = "abstract"
            elif tl == "references":
                section_name = "references"
            else:
                section_name = "body"

        # ── Headings (levels 1–3) ──
        if style == "Heading 1":
            format_heading(para, 1)
            prev_was_heading = True
            prev_was_label   = False
            continue

        if style == "Heading 2":
            format_heading(para, 2)
            prev_was_heading = True
            prev_was_label   = False
            continue

        if style == "Heading 3":
            format_heading(para, 3)
            prev_was_heading = True
            prev_was_label   = False
            continue

        # ── Skip empty paragraphs ──
        if not text:
            prev_was_heading = False
            prev_was_label   = False
            continue

        # ── Apply font + double spacing to all remaining text ──
        apply_font(para)
        set_double_spacing(para)

        pf = para.paragraph_format

        # ── References section: hanging indent ──
        if section_name == "references":
            pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent       = Inches(0.5)
            pf.first_line_indent = Inches(-0.5)

        # ── Abstract section: no first-line indent ──
        elif section_name == "abstract":
            pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent       = Inches(0)
            pf.first_line_indent = Pt(0)

        # ── Body ──
        else:
            pf.alignment   = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent = Inches(0)

            # Determine whether this paragraph should be indented.
            # No indent for:
            #   - First paragraph immediately after a heading
            #   - Table/Figure label lines ("Table 1", "Figure 3")
            #   - Table title lines (the line right after a label)
            #   - Note. lines beneath tables
            #   - Keywords: line
            is_label = bool(NO_INDENT_RE.match(text))
            no_indent = prev_was_heading or prev_was_label or is_label
            pf.first_line_indent = Pt(0) if no_indent else Inches(0.5)

            # Update label state for the NEXT paragraph
            # (catches the italic table title line that follows a label)
            prev_was_label = bool(
                re.match(r"^(Table|Figure)\s+\d+", text, re.IGNORECASE)
            )

        prev_was_heading = False

    # ── 3. Tables ─────────────────────────────────────────────────────────
    for table in doc.tables:
        format_apa7_table(table)

    # ── 4. Remove Pandoc bookmarks ────────────────────────────────────────
    remove_bookmarks(doc)

    # ── 5. Save ───────────────────────────────────────────────────────────
    doc.save(output_path)

    print(f"APA7 formatting applied -> {output_path}")
    print()
    print("Manual steps still required in Word:")
    print("  1. Page numbers  : Insert > Page Number > Top of Page > Plain Number 3")
    print("  2. Title page    : verify name, course, instructor, institution, date")
    print("  3. Figure labels : change italic 'Figure N' text to bold")
    print("  4. *Note.* lines : confirm only 'Note.' is italic, not the full sentence")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python3 apa7_format.py input.docx output.docx")
        sys.exit(1)

    format_document(sys.argv[1], sys.argv[2])
