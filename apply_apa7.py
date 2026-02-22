"""
APA7 Formatting Script for IoT Agriculture Final Report
Applies APA7 formatting to the Word document.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_PATH  = '/Users/dscottdawkins/usd/tiot/IoT_Agriculture_Final_Report.docx'
OUTPUT_PATH = '/Users/dscottdawkins/usd/tiot/IoT_Agriculture_Final_Report_APA7.docx'

FONT_NAME = 'Times New Roman'
FONT_SIZE = 12


# ---------------------------------------------------------------------------
# Helper: set font on a paragraph's runs AND its default run-properties
# ---------------------------------------------------------------------------
def set_para_font(para, font_name=FONT_NAME, font_size_pt=FONT_SIZE):
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)

    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)

    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(font_size_pt * 2))

    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(font_size_pt * 2))


# ---------------------------------------------------------------------------
# Helper: set spacing via direct XML (ensures line/lineRule are written)
# ---------------------------------------------------------------------------
def set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0):
    """
    Set line spacing directly via XML.
    double=True  => w:line=480, w:lineRule=auto  (double spacing)
    double=False => w:line=240, w:lineRule=auto  (single spacing)
    """
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)

    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)

    line_val = '480' if double else '240'
    spacing.set(qn('w:line'), line_val)
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), str(int(space_before_pt * 20)))
    spacing.set(qn('w:after'),  str(int(space_after_pt  * 20)))


# ---------------------------------------------------------------------------
# Helper: set first-line indent via XML
# ---------------------------------------------------------------------------
def set_indent_xml(para, first_line_twips=None, left_twips=None):
    """
    Set indentation directly via XML.
    Twips = 1/1440 inch. 1 inch = 1440 twips. 0.5 inch = 720 twips.
    Negative first_line_twips = hanging indent.
    """
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)

    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)

    # Remove existing ind attributes first
    for attr in list(ind.attrib.keys()):
        del ind.attrib[attr]

    if left_twips is not None:
        ind.set(qn('w:left'), str(int(left_twips)))
    if first_line_twips is not None:
        if first_line_twips < 0:
            # Hanging indent: use w:hanging (positive value)
            ind.set(qn('w:hanging'), str(int(abs(first_line_twips))))
        else:
            ind.set(qn('w:firstLine'), str(int(first_line_twips)))


def clear_indent_xml(para):
    """Remove all indentation from paragraph."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)


# ---------------------------------------------------------------------------
# Helper: clear shading / background color from a run
# ---------------------------------------------------------------------------
def clear_run_shading(run):
    rPr = run._r.find(qn('w:rPr'))
    if rPr is not None:
        for tag in ['w:shd', 'w:highlight']:
            elem = rPr.find(qn(tag))
            if elem is not None:
                rPr.remove(elem)


# ---------------------------------------------------------------------------
# Helper: clear shading from paragraph-level rPr
# ---------------------------------------------------------------------------
def clear_para_shading(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        rPr = pPr.find(qn('w:rPr'))
        if rPr is not None:
            shd2 = rPr.find(qn('w:shd'))
            if shd2 is not None:
                rPr.remove(shd2)


# ---------------------------------------------------------------------------
# Helper: detect heading style
# ---------------------------------------------------------------------------
def is_heading_style(para):
    if para.style and para.style.name:
        return 'heading' in para.style.name.lower()
    return False


# ---------------------------------------------------------------------------
# Table borders — APA7
# ---------------------------------------------------------------------------
def set_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')

    # Top border 1.5 pt (sz=12 = 12 * 1/8 pt = 1.5 pt)
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)

    # No left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'none')
    left.set(qn('w:sz'), '0')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), 'auto')
    tblBorders.append(left)

    # Bottom border 1.5 pt
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)

    # No right border
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'none')
    right.set(qn('w:sz'), '0')
    right.set(qn('w:space'), '0')
    right.set(qn('w:color'), 'auto')
    tblBorders.append(right)

    # No insideH
    insH = OxmlElement('w:insideH')
    insH.set(qn('w:val'), 'none')
    insH.set(qn('w:sz'), '0')
    insH.set(qn('w:space'), '0')
    insH.set(qn('w:color'), 'auto')
    tblBorders.append(insH)

    # No insideV
    insV = OxmlElement('w:insideV')
    insV.set(qn('w:val'), 'none')
    insV.set(qn('w:sz'), '0')
    insV.set(qn('w:space'), '0')
    insV.set(qn('w:color'), 'auto')
    tblBorders.append(insV)

    tblPr.append(tblBorders)

    # Header row: bottom border 1 pt (sz=8), clear all other borders
    if table.rows:
        first_row = table.rows[0]
        for cell in first_row.cells:
            tc   = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            existing_borders = tcPr.find(qn('w:tcBorders'))
            if existing_borders is not None:
                tcPr.remove(existing_borders)

            tcBorders = OxmlElement('w:tcBorders')

            bottom_b = OxmlElement('w:bottom')
            bottom_b.set(qn('w:val'), 'single')
            bottom_b.set(qn('w:sz'), '8')   # 1 pt
            bottom_b.set(qn('w:space'), '0')
            bottom_b.set(qn('w:color'), '000000')
            tcBorders.append(bottom_b)

            for side in ['top', 'left', 'right', 'insideH', 'insideV']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                b.set(qn('w:sz'), '0')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'auto')
                tcBorders.append(b)

            tcPr.append(tcBorders)

    # Non-header rows: clear all cell borders
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            continue
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            existing_borders = tcPr.find(qn('w:tcBorders'))
            if existing_borders is not None:
                tcPr.remove(existing_borders)
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                b.set(qn('w:sz'), '0')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'auto')
                tcBorders.append(b)
            tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# Table cell: clear background shading
# ---------------------------------------------------------------------------
def clear_cell_shading(cell):
    tc   = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'FFFFFF')


# ---------------------------------------------------------------------------
# Table cell padding
# ---------------------------------------------------------------------------
def set_table_cell_padding(table, top_pt=2, bottom_pt=2, left_pt=3, right_pt=3):
    def pt_to_twips(pt):
        return str(int(pt * 20))

    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)

    tblCellMar = OxmlElement('w:tblCellMar')
    for side, val in [('top', top_pt), ('left', left_pt), ('bottom', bottom_pt), ('right', right_pt)]:
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:w'), pt_to_twips(val))
        elem.set(qn('w:type'), 'dxa')
        tblCellMar.append(elem)
    tblPr.append(tblCellMar)


# ---------------------------------------------------------------------------
# Format cell text: Times New Roman 12pt, single-spaced, no extra spacing
# ---------------------------------------------------------------------------
def format_cell_text(cell, is_header=False):
    for para in cell.paragraphs:
        set_para_font(para, FONT_NAME, FONT_SIZE)
        # Single spacing for table cells
        set_spacing_xml(para, double=False, space_before_pt=0, space_after_pt=0)
        clear_indent_xml(para)

        clear_para_shading(para)

        for run in para.runs:
            clear_run_shading(run)
            if is_header:
                run.bold   = True
                run.italic = False
            run.font.color.rgb = RGBColor(0, 0, 0)

    clear_cell_shading(cell)


# ---------------------------------------------------------------------------
# Add page numbers to header (top-right)
# ---------------------------------------------------------------------------
def add_page_numbers(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    header = section.header
    # Clear existing header content
    for para in header.paragraphs:
        for run in para.runs:
            run.text = ''

    hdr_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_font(hdr_para, FONT_NAME, FONT_SIZE)

    pPr = hdr_para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        hdr_para._p.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'),  '0')

    run = hdr_para.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' PAGE '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Detect if a paragraph is in the References section
# ---------------------------------------------------------------------------
def build_references_index(paragraphs):
    for i, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        if is_heading_style(para) and 'references' in text:
            return i
    return None


# ---------------------------------------------------------------------------
# Find paragraphs preceding each table in the document body
# ---------------------------------------------------------------------------
def get_table_preceding_para_indices(doc):
    """
    Walk the document body children.
    Return {table_index: paragraph_index_before_table}
    """
    body      = doc.element.body
    para_idx  = -1
    table_idx = -1
    result    = {}
    last_para_idx = -1

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para_idx += 1
            last_para_idx = para_idx
        elif tag == 'tbl':
            table_idx += 1
            result[table_idx] = last_para_idx

    return result


# ---------------------------------------------------------------------------
# Format table label / title paragraphs
# ---------------------------------------------------------------------------
def format_table_label_paragraphs(paragraphs, table_prec_idx):
    for tbl_idx, prec_idx in table_prec_idx.items():
        if prec_idx < 0 or prec_idx >= len(paragraphs):
            continue
        prec_para = paragraphs[prec_idx]
        text = prec_para.text.strip()

        if re.match(r'^Table\s+\d+', text, re.IGNORECASE):
            # prec_para is the "Table N" label line — bold
            _apply_table_label_format(prec_para, bold=True, italic=False)
        elif prec_idx >= 1:
            prev_para = paragraphs[prec_idx - 1]
            prev_text = prev_para.text.strip()
            if re.match(r'^Table\s+\d+', prev_text, re.IGNORECASE):
                # prev_para = "Table N" label, prec_para = title
                _apply_table_label_format(prev_para, bold=True, italic=False)
                _apply_table_label_format(prec_para, bold=False, italic=True)


def _apply_table_label_format(para, bold, italic):
    set_para_font(para, FONT_NAME, FONT_SIZE)
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    clear_indent_xml(para)
    para.paragraph_format.left_indent = None
    for run in para.runs:
        run.bold   = bold
        run.italic = italic
        clear_run_shading(run)
    clear_para_shading(para)


# ---------------------------------------------------------------------------
# Format Heading 1 (paper title)
# ---------------------------------------------------------------------------
def format_heading1(para):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(para, FONT_NAME, FONT_SIZE)
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    clear_indent_xml(para)
    for run in para.runs:
        run.bold      = True
        run.italic    = False
        run.underline = False
        clear_run_shading(run)
    clear_para_shading(para)


# ---------------------------------------------------------------------------
# Format Heading 2 — APA7 Level 1 (centered, bold)
# ---------------------------------------------------------------------------
def format_heading2(para):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(para, FONT_NAME, FONT_SIZE)
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    clear_indent_xml(para)
    for run in para.runs:
        run.bold      = True
        run.italic    = False
        run.underline = False
        clear_run_shading(run)
    clear_para_shading(para)


# ---------------------------------------------------------------------------
# Format Heading 3 — APA7 Level 2 (flush left, bold)
# ---------------------------------------------------------------------------
def format_heading3(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_font(para, FONT_NAME, FONT_SIZE)
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    clear_indent_xml(para)
    for run in para.runs:
        run.bold      = True
        run.italic    = False
        run.underline = False
        clear_run_shading(run)
    clear_para_shading(para)


# ---------------------------------------------------------------------------
# Format Heading 4 — APA7 Level 3 (flush left, bold italic)
# ---------------------------------------------------------------------------
def format_heading4(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_font(para, FONT_NAME, FONT_SIZE)
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    clear_indent_xml(para)
    for run in para.runs:
        run.bold      = True
        run.italic    = True
        run.underline = False
        clear_run_shading(run)
    clear_para_shading(para)


# ---------------------------------------------------------------------------
# Format Keywords paragraph
# ---------------------------------------------------------------------------
def format_keywords_para(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    clear_indent_xml(para)

    full_text = para.text

    # Clear all existing runs
    for run in para.runs:
        run.text   = ''
        run.bold   = False
        run.italic = False
        clear_run_shading(run)

    if para.runs:
        kw_run = para.runs[0]
    else:
        kw_run = para.add_run()

    if 'Keywords:' in full_text:
        parts  = full_text.split('Keywords:', 1)
        prefix = parts[0]
        rest   = parts[1]

        kw_run.text      = prefix + 'Keywords:'
        kw_run.bold      = True
        kw_run.italic    = True
        kw_run.font.name = FONT_NAME
        kw_run.font.size = Pt(FONT_SIZE)

        rest_run           = para.add_run(rest)
        rest_run.bold      = False
        rest_run.italic    = True
        rest_run.font.name = FONT_NAME
        rest_run.font.size = Pt(FONT_SIZE)
    else:
        kw_run.text      = full_text
        kw_run.italic    = True
        kw_run.font.name = FONT_NAME
        kw_run.font.size = Pt(FONT_SIZE)

    # Set paragraph-level font too
    set_para_font(para, FONT_NAME, FONT_SIZE)
    clear_para_shading(para)


# ---------------------------------------------------------------------------
# Format a reference paragraph (hanging indent)
# ---------------------------------------------------------------------------
def format_reference_para(para):
    set_para_font(para, FONT_NAME, FONT_SIZE)
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    # Hanging indent: left=720 twips (0.5in), hanging=720 twips
    set_indent_xml(para, first_line_twips=-720, left_twips=720)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        clear_run_shading(run)
    clear_para_shading(para)


# ---------------------------------------------------------------------------
# Format a regular body paragraph
# ---------------------------------------------------------------------------
def format_body_para(para, first_line_indent=True):
    set_para_font(para, FONT_NAME, FONT_SIZE)
    set_spacing_xml(para, double=True, space_before_pt=0, space_after_pt=0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if first_line_indent:
        set_indent_xml(para, first_line_twips=720, left_twips=None)
    else:
        clear_indent_xml(para)
    for run in para.runs:
        clear_run_shading(run)
    clear_para_shading(para)


# ===========================================================================
# MAIN
# ===========================================================================
def main():
    print(f"Loading document: {INPUT_PATH}")
    doc = Document(INPUT_PATH)

    paragraphs = doc.paragraphs
    print(f"  Paragraphs: {len(paragraphs)}")
    print(f"  Tables:     {len(doc.tables)}")

    # --- Identify special paragraph indices ---
    abstract_idx = None
    keywords_idx = None

    for i, para in enumerate(paragraphs):
        text  = para.text.strip()
        tl    = text.lower()
        sname = para.style.name.lower() if para.style and para.style.name else ''

        if 'abstract' in tl and 'heading' in sname:
            if abstract_idx is None and i + 1 < len(paragraphs):
                abstract_idx = i + 1

        if tl.startswith('keywords'):
            keywords_idx = i

    refs_start = build_references_index(paragraphs)

    print(f"  Abstract para idx:  {abstract_idx}")
    print(f"  Keywords para idx:  {keywords_idx}")
    print(f"  References start:   {refs_start}")

    table_prec_idx = get_table_preceding_para_indices(doc)
    print(f"  Table preceding para indices: {table_prec_idx}")

    # Build set of para indices that are table labels or titles
    table_label_indices = set()
    for tbl_idx, prec_idx in table_prec_idx.items():
        if prec_idx >= 0:
            table_label_indices.add(prec_idx)
            if prec_idx >= 1:
                table_label_indices.add(prec_idx - 1)

    # --- Process all paragraphs ---
    print("Formatting paragraphs...")
    for i, para in enumerate(paragraphs):
        style_name  = para.style.name if para.style and para.style.name else ''
        sname_lower = style_name.lower()

        if 'heading 1' in sname_lower:
            format_heading1(para)

        elif 'heading 2' in sname_lower:
            format_heading2(para)

        elif 'heading 3' in sname_lower:
            format_heading3(para)

        elif 'heading 4' in sname_lower:
            format_heading4(para)

        elif i == keywords_idx:
            format_keywords_para(para)

        elif refs_start is not None and i > refs_start and para.text.strip():
            format_reference_para(para)

        else:
            no_indent = (
                i == 0 or
                i == abstract_idx or
                i == keywords_idx or
                i in table_label_indices or
                not para.text.strip() or
                para.text.strip().startswith('Note.') or
                para.text.strip().startswith('Note:')
            )
            format_body_para(para, first_line_indent=not no_indent)

    # --- Format table label paragraphs (post-pass to override body formatting) ---
    print("Formatting table label paragraphs...")
    format_table_label_paragraphs(paragraphs, table_prec_idx)

    # --- Format tables ---
    print(f"Formatting {len(doc.tables)} tables...")
    for t_idx, table in enumerate(doc.tables):
        print(f"  Table {t_idx + 1}...")
        set_table_borders(table)
        set_table_cell_padding(table, top_pt=2, bottom_pt=2, left_pt=3, right_pt=3)
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                format_cell_text(cell, is_header=is_header)

    # --- Add page numbers ---
    print("Adding page numbers to header...")
    add_page_numbers(doc)

    # --- Save ---
    print(f"Saving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!")


if __name__ == '__main__':
    main()
