#!/usr/bin/env python3
"""
create_apa7_reference.py
========================
Generates a reference.docx with APA7-compliant styles baked into every
Word style that Pandoc maps to.  Run this once, then pass the output to
every Pandoc conversion via --reference-doc.

Usage:
    python3 create_apa7_reference.py            # creates reference.docx
    python3 create_apa7_reference.py myref.docx # custom output path

Full pipeline:
    python3 create_apa7_reference.py
    pandoc paper.md --reference-doc=reference.docx -o paper_raw.docx
    python3 apa7_format.py paper_raw.docx paper_final.docx

Styles configured
-----------------
Paragraph styles (what Pandoc maps markdown elements to):
  Normal          – 12pt TNR, double-spaced, 0.5" first-line indent
  Body Text       – same as Normal (Pandoc fallback for body paragraphs)
  First Paragraph – same font/spacing, no first-line indent
  Compact         – same as Normal (Pandoc tight / bullet lists)
  Block Text      – 12pt TNR, double-spaced, 0.5" left+right indent (blockquotes)
  Heading 1       – 12pt TNR bold,        centered,      no indent
  Heading 2       – 12pt TNR bold,        left-aligned,  no indent
  Heading 3       – 12pt TNR bold-italic, left-aligned,  no indent
  Heading 4       – 12pt TNR italic,      left-aligned,  0.5" first-line
  Source Code     – 10pt Courier New, single-spaced,  0.5" left indent
  Caption         – 12pt TNR, left-aligned, no first-line indent
  Bibliography    – 12pt TNR, double-spaced, hanging indent (reference entries)

Character styles:
  Verbatim Char   – 10pt Courier New (inline code)

Page layout:
  US Letter, 1-inch margins all sides
"""

import sys
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

BODY_FONT   = "Times New Roman"
CODE_FONT   = "Courier New"
BODY_SIZE   = 12   # pt
CODE_SIZE   = 10   # pt
BLACK       = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _force_black(style):
    """
    Set font color to solid black and strip any theme color attributes
    (Word's built-in Heading styles default to a blue theme color that
    overrides explicit rgb settings unless cleared at the XML level).
    """
    # Set via the python-docx API first
    style.font.color.rgb = BLACK

    # Then reach into the XML and remove theme color attributes
    rPr = style._element.find(qn("w:rPr"))
    if rPr is None:
        return
    color_el = rPr.find(qn("w:color"))
    if color_el is None:
        return
    for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
        if qn(attr) in color_el.attrib:
            del color_el.attrib[qn(attr)]
    color_el.set(qn("w:val"), "000000")


def _remove_style_border(style):
    """
    Some built-in heading styles carry a bottom border.  Remove it so
    headings look clean in an APA7 document.
    """
    pPr = style._element.find(qn("w:pPr"))
    if pPr is None:
        return
    pb = pPr.find(qn("w:pBdr"))
    if pb is not None:
        pPr.remove(pb)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _get_or_create(doc, name, base_name=None, kind=WD_STYLE_TYPE.PARAGRAPH):
    """Return an existing style or create a new one of the given type."""
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, kind)
        if base_name:
            try:
                style.base_style = doc.styles[base_name]
            except KeyError:
                pass
        return style


def _apply_font(style, name=BODY_FONT, size=BODY_SIZE,
                bold=False, italic=False):
    """Apply font properties to a style and force black color."""
    style.font.name  = name
    style.font.size  = Pt(size)
    style.font.bold  = bold
    style.font.italic = italic
    _force_black(style)


def _apply_para_fmt(style,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_indent=Inches(0.5),
                    left_indent=Inches(0),
                    right_indent=Inches(0),
                    space_before=Pt(0),
                    space_after=Pt(0),
                    double_spaced=True):
    """Apply paragraph format properties to a style."""
    pf = style.paragraph_format
    pf.alignment         = alignment
    pf.first_line_indent = first_indent
    pf.left_indent       = left_indent
    pf.right_indent      = right_indent
    pf.space_before      = space_before
    pf.space_after       = space_after
    if double_spaced:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        # SINGLE doesn't write XML; use MULTIPLE @ 1.0 to force single spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.0


# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------

def configure_styles(doc):
    """Configure every style Pandoc may map to, according to APA7 rules."""

    # ── Normal ────────────────────────────────────────────────────────────
    # The base style for all body text
    normal = doc.styles["Normal"]
    _apply_font(normal)
    _apply_para_fmt(normal, first_indent=Inches(0.5))

    # ── Body Text ─────────────────────────────────────────────────────────
    # Pandoc sometimes emits Body Text instead of Normal
    body_text = _get_or_create(doc, "Body Text", base_name="Normal")
    _apply_font(body_text)
    _apply_para_fmt(body_text, first_indent=Inches(0.5))

    # ── First Paragraph ───────────────────────────────────────────────────
    # Pandoc uses this for the first paragraph after a heading (no indent)
    first_para = _get_or_create(doc, "First Paragraph", base_name="Normal")
    _apply_font(first_para)
    _apply_para_fmt(first_para, first_indent=Inches(0))

    # ── Compact ───────────────────────────────────────────────────────────
    # Pandoc uses Compact for tight (no blank line) lists
    compact = _get_or_create(doc, "Compact", base_name="Normal")
    _apply_font(compact)
    _apply_para_fmt(compact, first_indent=Inches(0))

    # ── Block Text ────────────────────────────────────────────────────────
    # Pandoc maps markdown block quotes to Block Text
    block = _get_or_create(doc, "Block Text", base_name="Normal")
    _apply_font(block)
    _apply_para_fmt(block,
                    first_indent=Inches(0),
                    left_indent=Inches(0.5),
                    right_indent=Inches(0.5))

    # ── Heading 1: centered, bold ─────────────────────────────────────────
    h1 = doc.styles["Heading 1"]
    _apply_font(h1, bold=True)
    _apply_para_fmt(h1,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_indent=Inches(0))
    _remove_style_border(h1)

    # ── Heading 2: left-aligned, bold ─────────────────────────────────────
    h2 = doc.styles["Heading 2"]
    _apply_font(h2, bold=True)
    _apply_para_fmt(h2, first_indent=Inches(0))
    _remove_style_border(h2)

    # ── Heading 3: left-aligned, bold italic ──────────────────────────────
    h3 = doc.styles["Heading 3"]
    _apply_font(h3, bold=True, italic=True)
    _apply_para_fmt(h3, first_indent=Inches(0))
    _remove_style_border(h3)

    # ── Heading 4: left-aligned, italic, indented (APA7 level 4) ──────────
    h4 = doc.styles["Heading 4"]
    _apply_font(h4, italic=True)
    _apply_para_fmt(h4, first_indent=Inches(0.5))
    _remove_style_border(h4)

    # ── Source Code ───────────────────────────────────────────────────────
    # Pandoc maps fenced code blocks to Source Code
    source_code = _get_or_create(doc, "Source Code", base_name="Normal")
    _apply_font(source_code, name=CODE_FONT, size=CODE_SIZE)
    _apply_para_fmt(source_code,
                    first_indent=Inches(0),
                    left_indent=Inches(0.5),
                    double_spaced=False)

    # ── Verbatim Char ─────────────────────────────────────────────────────
    # Pandoc maps inline backtick code to Verbatim Char (character style)
    verbatim = _get_or_create(doc, "Verbatim Char",
                               kind=WD_STYLE_TYPE.CHARACTER)
    verbatim.font.name = CODE_FONT
    verbatim.font.size = Pt(CODE_SIZE)

    # ── Caption ───────────────────────────────────────────────────────────
    # Figure and table captions
    caption = _get_or_create(doc, "Caption", base_name="Normal")
    _apply_font(caption)
    _apply_para_fmt(caption, first_indent=Inches(0))

    # ── Bibliography ──────────────────────────────────────────────────────
    # Reference list entries: hanging indent
    biblio = _get_or_create(doc, "Bibliography", base_name="Normal")
    _apply_font(biblio)
    _apply_para_fmt(biblio,
                    first_indent=Inches(-0.5),
                    left_indent=Inches(0.5))


# ---------------------------------------------------------------------------
# Page layout
# ---------------------------------------------------------------------------

def configure_page_layout(doc):
    """US Letter paper, 1-inch margins on all sides."""
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)


# ---------------------------------------------------------------------------
# Sample content
# ---------------------------------------------------------------------------

def add_sample_content(doc):
    """
    Add one paragraph using each style so all styles appear in the document
    body.  Pandoc ignores this content but having styles 'used' in the
    document ensures they survive round-trip DOCX processing.
    """
    doc.add_heading("APA7 Reference Template", level=1)

    doc.add_heading("Abstract", level=1)
    p = doc.add_paragraph(
        "Abstract body text. No first-line indent. Double-spaced."
    )
    p.style = doc.styles["Normal"]
    p.paragraph_format.first_line_indent = Inches(0)

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Body paragraph with 0.5-inch first-line indent. "
        "Double-spaced. Times New Roman 12pt."
    )

    doc.add_heading("Method", level=1)
    doc.add_heading("Participants", level=2)
    doc.add_paragraph("Body paragraph under Heading 2.")

    doc.add_heading("Materials", level=3)
    doc.add_paragraph("Body paragraph under Heading 3.")

    # Block quote
    bq = doc.add_paragraph(
        "Block quote text indented 0.5 inch on both sides."
    )
    bq.style = doc.styles["Block Text"]

    # Code block
    code = doc.add_paragraph('sample_variable = "hello world"')
    code.style = doc.styles["Source Code"]

    # Caption
    cap = doc.add_paragraph("Figure 1")
    cap.style = doc.styles["Caption"]

    # References
    doc.add_heading("References", level=1)
    ref = doc.add_paragraph(
        "Author, A. A. (2024). Title of work: Subtitle. Publisher. "
        "https://doi.org/10.0000/example"
    )
    ref.style = doc.styles["Bibliography"]


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def create_reference_doc(output_path="reference.docx"):
    doc = Document()
    configure_page_layout(doc)
    configure_styles(doc)
    add_sample_content(doc)
    doc.save(output_path)

    print(f"APA7 reference document created -> {output_path}")
    print()
    print("─" * 60)
    print("Complete pipeline:")
    print()
    print("  # Step 1 – install pandoc (once)")
    print("  brew install pandoc")
    print()
    print("  # Step 2 – convert markdown to raw DOCX")
    print(f"  pandoc paper.md --reference-doc={output_path} -o paper_raw.docx")
    print()
    print("  # Step 3 – apply remaining APA7 rules")
    print("  python3 apa7_format.py paper_raw.docx paper_final.docx")
    print()
    print("  # Step 4 – open paper_final.docx in Word and:")
    print("    • Add page numbers (Insert > Page Number > Top Right)")
    print("    • Add/verify title page")
    print("─" * 60)


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "reference.docx"
    create_reference_doc(output)
